"""
Bu dosya, terminal üzerinden çalışan gelişmiş bir Quiz / Test uygulamasıdır.

Amaç:
- Soruları CSV dosyasından okumak
- Kullanıcıya soruları sırayla sormak
- Cevapları kontrol etmek
- Sonuçları TXT, CSV, HTML, Word, Excel ve XML formatında raporlamak
- Soruları rastgele karıştırarak daha dinamik bir deneyim sunmak

Not:
Bu sürümde mevcut çalışan yapı korunmuştur.
Sadece kodun anlaşılmasını kolaylaştırmak için komutların ve blokların üstüne
çok detaylı açıklamalar eklenmiştir.
"""

# csv modülü:
# CSV (Comma Separated Values) dosyalarını okumak ve yazmak için kullanılır.
# Bu projede soru havuzunu 'questions.csv' dosyasından okumak ve
# quiz sonuçlarını '.csv' formatında dışarı aktarmak için kullanılmaktadır.
import csv

# html modülü:
# Kullanıcıya gösterilecek metinlerde yer alan özel karakterlerin
# HTML içinde güvenli biçimde gösterilmesini sağlar.
# Örneğin '<', '>', '&' gibi karakterler HTML yapısını bozmasın diye escape edilir.
# Bu sayede HTML raporunda soru ve şık metinleri güvenli şekilde gösterilir.
import html

# xml.etree.ElementTree modülü:
# Python standart kütüphanesi içinde gelir.
# Ekstra kurulum gerektirmez.
# Quiz sonuçlarını .xml formatında dışarı aktarmak için kullanılır.
import xml.etree.ElementTree as ET

# xml.dom.minidom modülü:
# XML çıktısını okunabilir, girintili ve düzenli hale getirmek için kullanılır.
from xml.dom import minidom

# random modülü:
# Rastgelelik gerektiren işlemlerde kullanılır.
# Bu projede soru listesini her quiz başlangıcında karıştırmak için kullanılır.
# Böylece her çalıştırmada sorular farklı sırada gelebilir.
import random

# datetime sınıfı:
# Tarih ve saat bilgisini almak için kullanılır.
# Burada sonuç dosyalarının isimlerini zaman damgası ile üretmek,
# ayrıca raporlara quizin hangi tarihte oluşturulduğunu yazmak için kullanılır.
from datetime import datetime

# Path sınıfı:
# Dosya ve klasör yollarını platformdan bağımsız, güvenli ve okunaklı şekilde
# yönetmek için kullanılır.
# Windows, Linux, macOS gibi sistemlerde dosya yollarını elle birleştirmek yerine
# nesne tabanlı güvenli yaklaşım sağlar.
from pathlib import Path

# zipfile modülü:
# DOCX ve XLSX dosyaları teknik olarak XML dosyalarının ZIP paketi halinde saklanmış biçimidir.
# Normalde Word için python-docx, Excel için openpyxl kullanılır.
# Ancak bu paketler kurulu değilse bile temel Word/Excel dosyası üretmek için standart kütüphane desteği sağlar.
import zipfile

# txt, csv, html, docx
# python-docx kütüphanesi
# Microsoft Word  .docx
# pip install python-docx
# Not:
# Bu importlar try/except içinde alınmıştır.
# Amaç: python-docx veya lxml kurulu değilse programın tamamen çökmesini engellemek.
# Paketler doğru kurulursa Word çıktısı normal şekilde üretilecektir.
try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ModuleNotFoundError:
    Document = None
    WD_TABLE_ALIGNMENT = None
    WD_ALIGN_PARAGRAPH = None
    Pt = None
    RGBColor = None
except ImportError:
    Document = None
    WD_TABLE_ALIGNMENT = None
    WD_ALIGN_PARAGRAPH = None
    Pt = None
    RGBColor = None

# txt, csv, html, docx, xlsx, .xml
# openpyxl kütüphanesi (Excel)
# Microsoft Excel  .xlsx
# pip install openpyxl
# Not:
# Bu importlar try/except içinde alınmıştır.
# Amaç: openpyxl kurulu değilse TXT, CSV, HTML ve XML çıktılarının yine de üretilmesidir.
try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
except ModuleNotFoundError:
    Workbook = None
    Alignment = None
    Border = None
    Font = None
    PatternFill = None
    Side = None
    get_column_letter = None
except ImportError:
    Workbook = None
    Alignment = None
    Border = None
    Font = None
    PatternFill = None
    Side = None
    get_column_letter = None

# BASE_DIR:
# Bu Python dosyasının bulunduğu klasörü temsil eder.
# __file__ -> o anda çalışan dosyanın yolunu verir.
# resolve() -> tam/gerçek yolu çözer.
# parent -> dosyanın bulunduğu klasörü alır.
# Amaç: CSV ve sonuç klasörlerini bu dosyanın bulunduğu dizine göre oluşturmak.
BASE_DIR = Path(__file__).resolve().parent

# CSV_FILE:
# Soruların okunacağı questions.csv dosyasının yolunu temsil eder.
# BASE_DIR / "questions.csv" ifadesi, işletim sistemine uygun biçimde
# dosya yolunu birleştirir.
CSV_FILE = BASE_DIR / "questions.csv"

# RESULTS_DIR:
# Quiz sonuçlarının kaydedileceği klasörü temsil eder.
# TXT, CSV ve HTML raporları bu klasör altına yazılacaktır.
RESULTS_DIR = BASE_DIR / "results"


# ensure_results_dir fonksiyonu:
# Sonuçların yazılacağı klasörün var olup olmadığını kontrol eder.
# Eğer klasör yoksa otomatik olarak oluşturur.
# parents=True  -> gerekirse üst klasörleri de oluşturur.
# exist_ok=True -> klasör zaten varsa hata vermez.
def ensure_results_dir():
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)


# load_questions fonksiyonu:
# Verilen CSV dosyasını okuyarak quiz sorularını belleğe yükler.
# Geriye liste döner.
#
# Beklenen kolonlar:
# - question
# - option_a
# - option_b
# - option_c
# - option_d
# - answer
#
# Bu fonksiyon ayrıca veri doğrulama da yapar:
# - Dosya var mı?
# - Başlıklar okunabiliyor mu?
# - Zorunlu kolonlar eksik mi?
# - Cevap A/B/C/D dışında mı?
# - Soru metni boş mu?
def load_questions(csv_file):
    # questions listesi, geçerli bulunan tüm soruları tutar.
    questions = []

    # Eğer CSV dosyası yoksa kullanıcıya hata gösterilir ve boş liste döndürülür.
    if not csv_file.exists():
        print(f"Hata: CSV dosyası bulunamadı -> {csv_file}")
        return questions

    # CSV dosyasını UTF-8-SIG ile açıyoruz.
    # utf-8-sig seçiminin nedeni:
    # Bazı CSV dosyaları başında BOM karakteri içerir.
    # Bu encoding, başlıkların bozulmadan okunmasına yardımcı olur.
    # newline="" parametresi, CSV okuma/yazma sırasında satır sonu sorunlarını azaltır.
    with open(csv_file, "r", encoding="utf-8-sig", newline="") as file:
        # DictReader, her satırı sözlük (dict) olarak okur.
        # Böylece sütunlara index ile değil isimleriyle erişebiliriz.
        reader = csv.DictReader(file)

        # required_columns kümesi:
        # CSV içinde mutlaka bulunması gereken başlıklar burada tanımlanır.
        required_columns = {
            "question",
            "option_a",
            "option_b",
            "option_c",
            "option_d",
            "answer",
        }

        # fieldnames boşsa başlıklar okunamamış demektir.
        if not reader.fieldnames:
            print("Hata: CSV başlıkları okunamadı.")
            return questions

        # missing hesabı:
        # Beklenen kolonlar ile gerçek kolonları karşılaştırır.
        # CSV'de eksik olan kolonları bulur.
        missing = required_columns - set(reader.fieldnames)
        if missing:
            print("Hata: CSV içinde eksik kolonlar var:", ", ".join(sorted(missing)))
            return questions

        # enumerate(reader, start=2):
        # CSV'de veri satırlarını gezer.
        # start=2 kullanılmasının nedeni:
        # 1. satır başlık satırıdır; gerçek veri satırları 2. satırdan başlar.
        for row_number, row in enumerate(reader, start=2):
            # answer alanı metne çevrilir, boşlukları temizlenir ve büyük harfe dönüştürülür.
            # Amaç: kullanıcı veri girişindeki küçük/büyük harf farklarını normalize etmek.
            answer = str(row["answer"]).strip().upper()

            # Eğer cevap yalnızca A/B/C/D dışında ise satır geçersiz sayılır.
            if answer not in {"A", "B", "C", "D"}:
                print(f"Uyarı: {row_number}. satırdaki cevap geçersiz olduğu için soru atlandı.")
                continue

            # Soru metni boşluklardan temizlenir.
            question_text = str(row["question"]).strip()

            # Şıklar sözlük halinde tutulur.
            # Böylece hem ekrana yazdırmak hem de sonradan karşılaştırmak kolaylaşır.
            options = {
                "A": str(row["option_a"]).strip(),
                "B": str(row["option_b"]).strip(),
                "C": str(row["option_c"]).strip(),
                "D": str(row["option_d"]).strip(),
            }

            # Soru metni boşsa kullanıcıya uyarı verilir ve bu satır alınmaz.
            if not question_text:
                print(f"Uyarı: {row_number}. satırdaki soru boş olduğu için atlandı.")
                continue

            # Geçerli bulunan soru, standart bir veri yapısı ile listeye eklenir.
            questions.append(
                {
                    "question": question_text,
                    "options": options,
                    "answer": answer,
                }
            )

    # Tüm geçerli sorular yüklendikten sonra geri döndürülür.
    return questions


# save_results_html fonksiyonu:
# Sonuçları görsel açıdan daha zengin bir HTML raporu olarak kaydeder.
# Kullanıcı hangi şıkkı işaretledi, doğru cevap neydi, hangi soru doğru/yanlıştı
# gibi bilgiler renkli kutular ve rozetlerle gösterilir.
def save_results_html(base_name, score, total, percent, user_results):
    html_path = base_name.with_suffix(".html")

    # Her soru kartını HTML içinde bir bölüm olarak biriktireceğiz.
    sections = []

    # Tüm kullanıcı sonuçlarını geziyoruz.
    for index, item in enumerate(user_results, start=1):
        options_html = []

        # Her şık için ayrı HTML satırı üretilir.
        for key, value in item["options"].items():
            # Şıkkın hangi görsel sınıfa sahip olacağı hesaplanır.
            css_class = get_option_css_class(key, item["user_answer"], item["correct_answer"])

            # html.escape ile içerik güvenli hale getirilir.
            label_parts = [f"<strong>{key})</strong> {html.escape(value)}"]

            # Kullanıcının seçtiği şık ise rozet eklenir.
            if key == item["user_answer"]:
                label_parts.append('<span class="badge badge-user">İşaretlediğin</span>')

            # Doğru cevap olan şık ise rozet eklenir.
            if key == item["correct_answer"]:
                label_parts.append('<span class="badge badge-correct">Doğru Cevap</span>')

            # Tek bir şık satırı HTML listesi elemanı olarak oluşturulur.
            option_line = f'<li class="{css_class}">{" ".join(label_parts)}</li>'
            options_html.append(option_line)

        # Soru doğru mu yanlış mı bilgisi hem metin hem CSS sınıfı olarak hazırlanır.
        question_status = "Doğru" if item["is_correct"] else "Yanlış"
        question_status_class = "status-correct" if item["is_correct"] else "status-wrong"

        # Tek bir soru kartının HTML içeriği hazırlanır.
        section = f"""
        <div class="question-card">
            <div class="question-header">
                <h2>Soru {index}</h2>
                <span class="status {question_status_class}">{question_status}</span>
            </div>
            <p class="question-text">{html.escape(item["question"])}</p>
            <ul class="options">
                {''.join(options_html)}
            </ul>
            <div class="answer-summary">
                <p><strong>İşaretlediğin:</strong> {item["user_answer"]}) {html.escape(item["options"][item["user_answer"]])}</p>
                <p><strong>Doğru cevap:</strong> {item["correct_answer"]}) {html.escape(item["options"][item["correct_answer"]])}</p>
            </div>
        </div>
        """
        sections.append(section)

    # Tüm raporun ana HTML iskeleti hazırlanır.
    # İçinde CSS stilleri de gömülü olduğu için tek dosya halinde açılabilir.
    html_content = f"""<!DOCTYPE html>
<html lang="tr">
<head>
    <meta charset="UTF-8">
    <title>Quiz Sonuç Raporu</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            background: #f4f6f8;
            color: #1f2937;
            margin: 0;
            padding: 24px;
        }}
        .container {{
            max-width: 1000px;
            margin: 0 auto;
        }}
        .summary {{
            background: #ffffff;
            border-radius: 14px;
            padding: 20px;
            margin-bottom: 24px;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
        }}
        .summary h1 {{
            margin-top: 0;
        }}
        .summary-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 12px;
            margin-top: 16px;
        }}
        .summary-box {{
            background: #f9fafb;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 14px;
        }}
        .question-card {{
            background: #ffffff;
            border-radius: 14px;
            padding: 20px;
            margin-bottom: 18px;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
        }}
        .question-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            gap: 12px;
            flex-wrap: wrap;
        }}
        .question-text {{
            font-size: 18px;
            font-weight: bold;
        }}
        .status {{
            padding: 8px 12px;
            border-radius: 999px;
            font-weight: bold;
            font-size: 14px;
        }}
        .status-correct {{
            background: #dcfce7;
            color: #166534;
        }}
        .status-wrong {{
            background: #fee2e2;
            color: #991b1b;
        }}
        .options {{
            list-style: none;
            padding: 0;
            margin: 16px 0;
        }}
        .option {{
            background: #f9fafb;
            border: 1px solid #d1d5db;
            border-radius: 10px;
            padding: 12px;
            margin-bottom: 10px;
        }}
        .option-correct {{
            background: #dcfce7;
            border-color: #22c55e;
        }}
        .option-correct-selected {{
            background: #bbf7d0;
            border: 2px solid #15803d;
        }}
        .option-wrong-selected {{
            background: #fee2e2;
            border-color: #ef4444;
        }}
        .badge {{
            display: inline-block;
            margin-left: 8px;
            padding: 4px 8px;
            border-radius: 999px;
            font-size: 12px;
            font-weight: bold;
        }}
        .badge-user {{
            background: #dbeafe;
            color: #1d4ed8;
        }}
        .badge-correct {{
            background: #dcfce7;
            color: #166534;
        }}
        .answer-summary {{
            background: #f9fafb;
            border-radius: 10px;
            padding: 12px;
            border: 1px solid #e5e7eb;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="summary">
            <h1>Quiz Sonuç Raporu</h1>
            <p><strong>Tarih:</strong> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
            <div class="summary-grid">
                <div class="summary-box"><strong>Doğru Sayısı</strong><br>{score}</div>
                <div class="summary-box"><strong>Yanlış Sayısı</strong><br>{total - score}</div>
                <div class="summary-box"><strong>Toplam Soru</strong><br>{total}</div>
                <div class="summary-box"><strong>Başarı Oranı</strong><br>%{percent:.2f}</div>
            </div>
        </div>

        {''.join(sections)}
    </div>
</body>
</html>
"""

    # Hazırlanan HTML metni dosyaya UTF-8 ile yazılır.
    html_path.write_text(html_content, encoding="utf-8")
    return html_path


# prompt_menu fonksiyonu:
# Kullanıcıya ana menüyü gösterir ve seçim alır.
# Geçersiz giriş yapılırsa kullanıcıyı yeniden yönlendirir.
# Sadece 1, 2 veya 3 kabul edilir.
def prompt_menu():
    print("\n" + "=" * 70)
    print("Python Quiz / Test Uygulaması")
    print("=" * 70)
    print("1) Quiz başlat")
    print("2) Soru sayısını göster")
    print("3) Çıkış")

    # input ile kullanıcıdan seçim alınır.
    # strip() ile baştaki/sondaki boşluklar temizlenir.
    choice = input("Seçiminiz: ").strip()

    # while döngüsü, geçerli bir menü seçimi yapılana kadar tekrar ister.
    while choice not in {"1", "2", "3"}:
        choice = input("Geçersiz seçim. Lütfen 1, 2 veya 3 girin: ").strip()
    return choice


# ask_question fonksiyonu:
# Tek bir soruyu kullanıcıya gösterir, cevabı alır ve doğru/yanlış kontrolü yapar.
#
# Parametreler:
# - index: Kaçıncı soru olduğu
# - total: Toplam soru sayısı
# - question_data: Soru metni, şıklar ve doğru cevabı içeren sözlük
#
# Geriye, bu soruya ait kullanıcı sonucunu içeren bir sözlük döndürür.
def ask_question(index, total, question_data):
    print("\n" + "=" * 70)
    print(f"Soru {index}/{total}")
    print("-" * 70)
    print(question_data["question"])

    # options sözlüğündeki tüm şıkları sırayla ekranda göstersin
    for key, value in question_data["options"].items():
        print(f"{key}) {value}")

    # Kullanıcıdan cevap alınır.
    # upper() kullanımı sayesinde küçük harf girilirse bile büyük harfe çevirsin.
    user_answer = input("\nCevabınız: (A/B/C/D): ").strip().upper()

    # Geçerli bir cevap girilene akdar kullancıı tekrar yönlendilir.
    while user_answer not in {"A", "B", "C", "D"}:
        user_answer = input("Geçersiz giriş. Lütfen A,B,D veya D giriniz: ").strip().upper()

    # Soru verisindeki doğru cevap alınır.
    correct_answer = question_data["answer"]

    # Kullanıcını cevabı ile dorğu cevap karşılaştırılır.
    is_correct = user_answer == correct_answer

    # Kullanıcı anlık geri bildirim verilir
    if is_correct:
        print("Sonuç: Doğru")
    else:
        print(
            f"Sonuç: Yanlış | Doğru cevap: {correct_answer} "
            f"{question_data['options'][correct_answer]}"
        )

    # Sonuçlar standart yapıda döndürülsün
    return {
        "question": question_data["question"],
        "options": question_data["options"],
        "correct_answer": correct_answer,
        "user_answer": user_answer,
        "is_correct": is_correct,
    }


# Quiz akışının ana  çalıştırıcı fonksiyonudur.
# Tüm soruları rastgele karıştırır ve sırayla kullanıcıya sorar
# puan hesaplar ve sonuç listesinini dönderir.
def run_quiz(questions):
    # Original soru listesini bozmamak için kopyalama yapıyoruz.
    randomized_questions = questions[:]

    # Sorular rastgele karıştırılır
    # Böylece her quiz aynı sırada başlamaz
    random.shuffle(randomized_questions)

    print("\nQuiz başladı. Sroular rastgele karıştırıldı")

    # Sorular rastgele sayısı hespalanır.
    total = len(randomized_questions)

    # Score değişkeni doğru cevap sayısını tutar
    score = 0

    # user_result listesi, her soruya ilişkin kullanıcı sonucunu saklar.
    user_results = []

    # enumarate : soru numaralandırılması
    for index, question_data in enumerate(randomized_questions, start=1):
        result = ask_question(index, total, question_data)
        user_results.append(result)

        # Eğer soru doğru ise skor 1 artırır
        if result["is_correct"]:
            # score =score+ 1
            score += 1

    # Quiz tamamlandıktan sonra puan, toplam soru ve detaylı sonuçları döndersin.
    return score, total, user_results


# create_result_base_names fonksiyonu:
# Sonuç dosyalarını ortak temel isi üretir.
# Örneğin:
# results/quiz_result_20260428_171845
# Sonra buna .txt, .csv, .html uzantıları eklenir
def create_result_base_names():
    # önce sonuç klasörünün var olup olmadığını kontrol edelim
    ensure_results_dir()

    # Şu anki tarih-saat bilgisi dosya adı için uygun formatta çevrilir
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Path nesnesi olarak temel dosya adını oluşturulur
    base_name = RESULTS_DIR / f"quiz_result_{timestamp}"
    return base_name


#############################################################################
# TXT CREATE
# save_results_txt dosyasını fonksiyonun
# Quiz sonucu okunabilir bir metin raporu olarak '.txt' dosyanına  kaydeder.
# Text insan gözüyle daha rahat okunur
def save_results_txt(base_name, score, total, percent, user_results):
    txt_path = base_name.with_suffix(".txt")

    # Dosya yazma modunda açılır
    with open(txt_path, "w", encoding="utf-8") as file:
        file.write("PYTHON QUIZ SONUÇ RAPORU\n")
        file.write("=" * 70 + "\n")
        file.write(f"Tarih          : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        file.write(f"Doğru sayısı   : {score}\n")
        file.write(f"Yanlış sayısı  : {total - score}\n")
        file.write(f"Toplam soru    : {total}\n")
        file.write(f"Başarı Oranı   : {percent:.2f}\n")
        file.write("=" * 70 + "\n\n")

        # Her soru tek tek rapora detaylı şekilde yazılır.
        for index, item in enumerate(user_results, start=1):
            file.write(f"Soru   {index}: {item['question']}\n")
            file.write(f"İşaretlenen cevap :  {item['user_answer']}) {item['options'][item['user_answer']]}\n")
            file.write(f"Doğru cevap       :  {item['correct_answer']}) {item['options'][item['correct_answer']]}\n")
            file.write(f"Durum             :  {'Doğru' if item['is_correct'] else 'Yanlış'}\n")
            file.write("-" * 70 + "\n")

    return txt_path


#############################################################################
# CSV CREATE
# save_result_csv dosyasını fonksiyonun
# Sonuçları tablo yapısında '.csv' olarak kaydeder
# böylece Excel, Google Sheets veya veri analizi araçlarında kolaylıkla açılabilir
# Text insan gözüyle daha rahat okunur
# def save_results_csv
def save_results_csv(base_name, score, total, percent, user_results):
    csv_path = base_name.with_suffix(".csv")

    # Dosya yazma modunda açılır
    with open(csv_path, "w", encoding="utf-8") as file:
        writer = csv.writer(file)

        # özet bilgiler
        writer.writerow(["summary_type", "value"])
        writer.writerow(["date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
        writer.writerow(["score", score])
        writer.writerow(["wrong", total - score])
        writer.writerow(["total", total])
        writer.writerow(["percent", f"{percent:.2f}"])
        writer.writerow([])

        # Ardından detay veri başlıkları yazılır.
        writer.writerow([
            "question_no",
            "question",
            "option_a",
            "option_b",
            "option_c",
            "option_d",
            "user_answer",
            "correct_answer",
            "result",
        ])

        # Her sorunun detayları satır satır yazılır.
        for index, item in enumerate(user_results, start=1):
            writer.writerow([
                index,
                item["question"],
                item["options"]["A"],
                item["options"]["B"],
                item["options"]["C"],
                item["options"]["D"],
                item["user_answer"],
                item["correct_answer"],
                "Doğru" if item["is_correct"] else "Yanlış",
            ])

    return csv_path


#############################################################################
# DOCX(WORD) CREATE
# add_docx_paragraph fonksiyonu:
# Word raporunda tekrar eden paragraf oluşturma işlemini sadeleştirir
# run.bold ve run.font.size gibi değerler burada merkezi şekilde yönetilir.
def add_docx_paragraph(document, text, bold=False, font_size=11, color=None):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)

    # Renk verilmişse RGBColor ile uygulanır.
    if color:
        run.font.color.rgb = RGBColor(*color)

    return paragraph



# save_results_docx_fallback fonksiyonu:
# python-docx veya onun arka plandaki lxml bağımlılığı kurulu değilse devreye girer.
# DOCX dosyası, Office Open XML standardına göre ZIP içindeki XML dosyalarından oluşur.
# Bu fonksiyon sade ama Word tarafından açılabilen temel bir .docx raporu üretir.
# Amaç: Harici paket kurulmasa bile Word çıktısının tamamen kaybolmasını engellemektir.
def save_results_docx_fallback(base_name, score, total, percent, user_results):
    docx_path = base_name.with_suffix(".docx")

    # Word XML içine yazılacak özel karakterleri güvenli hale getirir.
    def word_escape(value):
        return html.escape("" if value is None else str(value), quote=False)

    # Tek bir Word paragrafını XML olarak oluşturur.
    def word_paragraph(text, bold=False, color=None, center=False):
        run_properties = []
        paragraph_properties = []

        if bold:
            run_properties.append("<w:b/>")

        if color:
            if isinstance(color, tuple):
                color_hex = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
            else:
                color_hex = str(color)
            run_properties.append(f'<w:color w:val="{color_hex}"/>')

        if center:
            paragraph_properties.append('<w:jc w:val="center"/>')

        ppr_xml = f"<w:pPr>{''.join(paragraph_properties)}</w:pPr>" if paragraph_properties else ""
        rpr_xml = f"<w:rPr>{''.join(run_properties)}</w:rPr>" if run_properties else ""

        # xml:space="preserve" ile baştaki/sondaki boşlukların korunması sağlanır.
        return f"<w:p>{ppr_xml}<w:r>{rpr_xml}<w:t xml:space=\"preserve\">{word_escape(text)}</w:t></w:r></w:p>"

    # Word raporu için paragraf listesi hazırlanır.
    paragraphs = []
    paragraphs.append(word_paragraph("Python Quiz Sonuç Raporu", bold=True, center=True))
    paragraphs.append(word_paragraph(f"Tarih: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", bold=True))
    paragraphs.append(word_paragraph("Genel Özet", bold=True))
    paragraphs.append(word_paragraph(f"Doğru Sayısı: {score}"))
    paragraphs.append(word_paragraph(f"Yanlış Sayısı: {total - score}"))
    paragraphs.append(word_paragraph(f"Toplam Soru: {total}"))
    paragraphs.append(word_paragraph(f"Başarı Oranı: %{percent:.2f}"))
    paragraphs.append(word_paragraph("Soru Detayları", bold=True))

    # Her soru için detaylı Word içeriği eklenir.
    for index, item in enumerate(user_results, start=1):
        status_text = "Doğru" if item["is_correct"] else "Yanlış"
        status_color = (22, 101, 52) if item["is_correct"] else (153, 27, 27)

        paragraphs.append(word_paragraph(f"Soru {index} - {status_text}", bold=True, color=status_color))
        paragraphs.append(word_paragraph(item["question"], bold=True))

        for option_key, option_value in item["options"].items():
            option_suffix = []
            if option_key == item["user_answer"]:
                option_suffix.append("İşaretlediğin")
            if option_key == item["correct_answer"]:
                option_suffix.append("Doğru Cevap")

            suffix_text = f"  [{' | '.join(option_suffix)}]" if option_suffix else ""
            paragraphs.append(word_paragraph(f"{option_key}) {option_value}{suffix_text}"))

        paragraphs.append(word_paragraph(f"İşaretlenen cevap: {item['user_answer']}) {item['options'][item['user_answer']]}"))
        paragraphs.append(word_paragraph(f"Doğru cevap: {item['correct_answer']}) {item['options'][item['correct_answer']]}"))
        paragraphs.append(word_paragraph("-" * 70))

    # Word ana doküman XML içeriği hazırlanır.
    document_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:body>
        {''.join(paragraphs)}
        <w:sectPr>
            <w:pgSz w:w="11906" w:h="16838"/>
            <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="708" w:footer="708" w:gutter="0"/>
        </w:sectPr>
    </w:body>
</w:document>'''

    # DOCX paketinin içerik tipleri tanımlanır.
    content_types_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
    <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
    <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>'''

    # DOCX ana ilişki dosyası hazırlanır.
    rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
    <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
    <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>'''

    # DOCX belge özellikleri hazırlanır.
    core_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
    <dc:title>Python Quiz Sonuç Raporu</dc:title>
    <dc:creator>Python Quiz Uygulaması</dc:creator>
    <cp:lastModifiedBy>Python Quiz Uygulaması</cp:lastModifiedBy>
    <dcterms:created xsi:type="dcterms:W3CDTF">{datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')}</dcterms:created>
    <dcterms:modified xsi:type="dcterms:W3CDTF">{datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')}</dcterms:modified>
</cp:coreProperties>'''

    app_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
    <Application>Python Quiz Uygulaması</Application>
</Properties>'''

    # Tüm XML parçaları .docx zip paketi içine yazılır.
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as docx_zip:
        docx_zip.writestr("[Content_Types].xml", content_types_xml)
        docx_zip.writestr("_rels/.rels", rels_xml)
        docx_zip.writestr("word/document.xml", document_xml)
        docx_zip.writestr("docProps/core.xml", core_xml)
        docx_zip.writestr("docProps/app.xml", app_xml)

    return docx_path


# save_results_docx fonksiyonu:
# Quiz sonucunu Microsoft Word tarafından açılabilen '.docx' formatından kaydetmek
# Bu raporda önce özet bilgiler, sonra her sorunun detaylı cevabı yer alır.
# Amaç: Kulalnıcı çıktıyı doğrudan rapor/doküman olarak paylaşabilsin
def save_results_docx(base_name, score, total, percent, user_results):
    docx_path = base_name.with_suffix(".docx")

    # python-docx kurulmamışsa standart Python fallback sistemi devreye girer.
    # Böylece Word çıktısı tamamen iptal edilmez, sade bir .docx raporu yine üretilir.
    if Document is None:
        print("Uyarı: python-docx/lxml bulunamadı. Standart Python fallback ile temel Word raporu oluşturuluyor.")
        print("Daha gelişmiş Word biçimlendirmesi için kurulum: python -m pip install python-docx lxml")
        return save_results_docx_fallback(base_name, score, total, percent, user_results)

    # Yeni bir word dokümanı oluşturmak
    document = Document()

    # Ana başlık hazırlanır
    title = document.add_heading("Python Quiz Sonuç raporu", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Rapor tarihi ve genel özet alanı hazırlanır
    add_docx_paragraph(document, f"Tarih: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", bold=True)
    add_docx_paragraph(document, "Genel Özet", bold=True, font_size=14)

    # Özet Tablo: doğru, yanlış, toplam, başarı oranı bilgilerini içerir.
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.style = "Table Grid"

    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Alan"
    header_cells[1].text = "Değer"

    summary_rows = [
        ("Doğru Sayısı", str(score)),
        ("Yanlış Sayısı", str(total - score)),
        ("Toplam Soru", str(total)),
        ("Başarı Oranı", f"%{percent:.2f}"),
    ]

    # Döngü olarak
    for label, value in summary_rows:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    #
    document.add_paragraph()
    add_docx_paragraph(document, "Soru detayları", bold=True, font_size=14)

    # Her soru için Word içinde ayrı bölüm oluşturulur.
    for index, item in enumerate(user_results, start=1):
        status_text = "Doğru" if item["is_correct"] else "Yanlış"
        status_color = (22, 101, 52) if item["is_correct"] else (153, 27, 27)

        add_docx_paragraph(
            document,
            f"Soru {index}- {status_text}",
            bold=True,
            font_size=12,
            color=status_color,
        )
        add_docx_paragraph(document, item["question"], bold=True)

        # Şıklar Word raporunda madde madde yazılır.
        for option_key, option_value in item["options"].items():
            option_suffix = []
            if option_key == item["user_answer"]:
                option_suffix.append("İşaretlediğin")
            if option_key == item["correct_answer"]:
                option_suffix.append("Doğru Cevap")

            suffix_text = f"  [{' | '.join(option_suffix)}]" if option_suffix else ""
            add_docx_paragraph(document, f"{option_key}) {option_value}{suffix_text}")

        add_docx_paragraph(
            document,
            f"İşaretlenen cevap: {item['user_answer']}) {item['options'][item['user_answer']]}",
        )

        add_docx_paragraph(
            document,
            f"Doğru cevap: {item['correct_answer']}) {item['options'][item['correct_answer']]}",
        )
        add_docx_paragraph(document, "-" * 70)

    # Word dosyası belirtilen yola kaydedilir.
    document.save(docx_path)
    return docx_path


################################################################################
# EXCEL CREATE
# set_excel_cell_style fonksiyonu
# Excel raporundaki başlık ve veri hücrelerinini daha okunabilir görünmesi için
# ortak font, kenarlık, hizalama, dolgu ayarlarını uygular
def set_excel_cell_style(cell, bold=False, fill_color=None, font_color="000000"):
    thin_border = Border(
        left=Side(style="thin", color="D1D5DB"),
        right=Side(style="thin", color="D1D5DB"),
        top=Side(style="thin", color="D1D5DB"),
        bottom=Side(style="thin", color="D1D5DB"),
    )

    cell.font = Font(bold=bold, color=font_color)
    cell.border = thin_border
    cell.alignment = Alignment(vertical="top", wrap_text=True)

    if fill_color:
        cell.fill = PatternFill(fill_type="solid", fgColor=fill_color)


# set_excel_style fonksiyonu:
# Excel rapoundaki başlık ve veri hüxrelerinin daha okunabilir görünmesi için
# ortak font, kenarlık, hizalama ve dolgu ayarlarını uygular
# Not:
# Bu fonksiyon önceki blokla aynı amaca hizmet eder.
# Yorumlar korunmuştur, ancak eski sürümdeki yazım hataları düzeltilmiştir.
def set_excel_cell_style(cell, bold=False, fill_color=None, font_color="000000"):
    thin_border = Border(
        left=Side(style="thin", color="D1D5DB"),
        right=Side(style="thin", color="D1D5DB"),
        top=Side(style="thin", color="D1D5DB"),
        bottom=Side(style="thin", color="D1D5DB"),
    )

    cell.font = Font(bold=bold, color=font_color)
    cell.border = thin_border
    cell.alignment = Alignment(vertical="top", wrap_text=True)

    if fill_color:
        cell.fill = PatternFill(fill_type="solid", fgColor=fill_color)


# autofit_excel_columns Fonksiyonu
# Excel sayfalarında kolon genişliğini içerik uzunluğuna göre ayarlar
# Çok uzun metinler için maksimum genişlik sınırı olması lazım.
def autofit_excel_columns(worksheet, max_width=45):
    for column_cells in worksheet.columns:
        max_length = 0
        column_letter = get_column_letter(column_cells[0].column)

        for cell in column_cells:
            value = cell.value
            if value is not None:
                max_length = max(max_length, len(str(value)))

        worksheet.column_dimensions[column_letter].width = min(max_length + 3, max_width)



# save_results_xlsx_fallback fonksiyonu:
# openpyxl kurulu değilse devreye girer.
# XLSX dosyası da DOCX gibi ZIP paketi içindeki XML dosyalarından oluşur.
# Bu fonksiyon sade ama Excel tarafından açılabilen temel bir .xlsx raporu üretir.
# Amaç: Harici paket kurulmasa bile Excel çıktısının tamamen kaybolmasını engellemektir.
def save_results_xlsx_fallback(base_name, score, total, percent, user_results):
    xlsx_path = base_name.with_suffix(".xlsx")

    # Excel hücre referansı için kolon numarasını harfe çevirir.
    # Örneğin: 1 -> A, 2 -> B, 27 -> AA
    def excel_column_letter(column_number):
        letters = ""
        while column_number:
            column_number, remainder = divmod(column_number - 1, 26)
            letters = chr(65 + remainder) + letters
        return letters

    # XML içine yazılacak özel karakterleri güvenli hale getirir.
    def excel_escape(value):
        return html.escape("" if value is None else str(value), quote=False)

    # Hücre XML içeriğini üretir.
    def excel_cell(row_number, column_number, value):
        cell_reference = f"{excel_column_letter(column_number)}{row_number}"

        # Sayısal değerlerde Excel'in sayı tipi kullanılabilir.
        if isinstance(value, (int, float)) and not isinstance(value, bool):
            return f'<c r="{cell_reference}"><v>{value}</v></c>'

        # Metin değerleri inline string olarak yazılır.
        return f'<c r="{cell_reference}" t="inlineStr"><is><t>{excel_escape(value)}</t></is></c>'

    # Satır XML içeriğini üretir.
    def excel_row(row_number, values):
        cells = []
        for column_number, value in enumerate(values, start=1):
            cells.append(excel_cell(row_number, column_number, value))
        return f'<row r="{row_number}">{"".join(cells)}</row>'

    # Sayfa XML içeriğini üretir.
    def worksheet_xml(rows):
        row_xml = []
        for row_number, values in enumerate(rows, start=1):
            row_xml.append(excel_row(row_number, values))

        return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
    <sheetData>
        {''.join(row_xml)}
    </sheetData>
</worksheet>'''

    # Özet sayfası verileri hazırlanır.
    summary_rows = [
        ["Python Quiz Sonuç Raporu"],
        [],
        ["Tarih", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
        ["Doğru Sayısı", score],
        ["Yanlış Sayısı", total - score],
        ["Toplam Soru", total],
        ["Başarı Oranı", f"%{percent:.2f}"],
    ]

    # Detay sayfası başlıkları hazırlanır.
    detail_rows = [[
        "Soru No",
        "Soru",
        "A Şıkkı",
        "B Şıkkı",
        "C Şıkkı",
        "D Şıkkı",
        "İşaretlenen Cevap",
        "Doğru Cevap",
        "Sonuç",
    ]]

    # Her soru sonucu detay sayfasına satır olarak eklenir.
    for index, item in enumerate(user_results, start=1):
        detail_rows.append([
            index,
            item["question"],
            item["options"]["A"],
            item["options"]["B"],
            item["options"]["C"],
            item["options"]["D"],
            item["user_answer"],
            item["correct_answer"],
            "Doğru" if item["is_correct"] else "Yanlış",
        ])

    # XLSX paketinin içerik tipleri tanımlanır.
    content_types_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
    <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
    <Override PartName="/xl/worksheets/sheet2.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
    <Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>
</Types>'''

    # XLSX ana ilişki dosyası hazırlanır.
    rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>'''

    # Workbook içinde iki sayfa tanımlanır: Özet ve Detaylar.
    workbook_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
    <sheets>
        <sheet name="Özet" sheetId="1" r:id="rId1"/>
        <sheet name="Detaylar" sheetId="2" r:id="rId2"/>
    </sheets>
</workbook>'''

    workbook_rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
    <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet2.xml"/>
    <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>'''

    # Temel stiller dosyası hazırlanır.
    styles_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
    <fonts count="1"><font><sz val="11"/><name val="Calibri"/></font></fonts>
    <fills count="1"><fill><patternFill patternType="none"/></fill></fills>
    <borders count="1"><border/></borders>
    <cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>
    <cellXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/></cellXfs>
</styleSheet>'''

    # Tüm XML parçaları .xlsx zip paketi içine yazılır.
    with zipfile.ZipFile(xlsx_path, "w", zipfile.ZIP_DEFLATED) as xlsx_zip:
        xlsx_zip.writestr("[Content_Types].xml", content_types_xml)
        xlsx_zip.writestr("_rels/.rels", rels_xml)
        xlsx_zip.writestr("xl/workbook.xml", workbook_xml)
        xlsx_zip.writestr("xl/_rels/workbook.xml.rels", workbook_rels_xml)
        xlsx_zip.writestr("xl/styles.xml", styles_xml)
        xlsx_zip.writestr("xl/worksheets/sheet1.xml", worksheet_xml(summary_rows))
        xlsx_zip.writestr("xl/worksheets/sheet2.xml", worksheet_xml(detail_rows))

    return xlsx_path


# save_results_xlsx fonksiyonu:
# Quiz sonucunu Microsoft Excel tarafından açılabilen '.xlsx' formatında kaydeder.
# Dosyada iki sayfa vardır:
# 1) Özet
# 2) Detaylar
# Amaç: sonuçların filtrelenebilir, tablolaştırılabilir ve analiz edilebilir olmasıdır.
def save_results_xlsx(base_name, score, total, percent, user_results):
    xlsx_path = base_name.with_suffix(".xlsx")

    # openpyxl kurulmamışsa standart Python fallback sistemi devreye girer.
    # Böylece Excel çıktısı tamamen iptal edilmez, sade bir .xlsx raporu yine üretilir.
    if Workbook is None:
        print("Uyarı: openpyxl bulunamadı. Standart Python fallback ile temel Excel raporu oluşturuluyor.")
        print("Daha gelişmiş Excel biçimlendirmesi için kurulum: python -m pip install openpyxl")
        return save_results_xlsx_fallback(base_name, score, total, percent, user_results)

    # Yeni Excel çalışma kitabı oluşturulur.
    workbook = Workbook()

    # Varsayılan sayfa özet sayfası olarak yeniden adlandırılır.
    summary_sheet = workbook.active
    summary_sheet.title = "Özet"

    # Detaylar için ikinci sayfa oluşturulur.
    details_sheet = workbook.create_sheet("Detaylar")

    # Özet sayfası başlığı.
    summary_sheet["A1"] = "Python Quiz Sonuç Raporu"
    summary_sheet["A1"].font = Font(bold=True, size=16, color="111827")
    summary_sheet.merge_cells("A1:B1")

    summary_data = [
        ("Tarih", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Doğru Sayısı", score),
        ("Yanlış Sayısı", total - score),
        ("Toplam Soru", total),
        ("Başarı Oranı", f"%{percent:.2f}"),
    ]

    # Özet bilgileri satır satır yazılır.
    for row_index, (label, value) in enumerate(summary_data, start=3):
        summary_sheet.cell(row=row_index, column=1, value=label)
        summary_sheet.cell(row=row_index, column=2, value=value)

    # Özet sayfasında stil uygulanır.
    for row in summary_sheet.iter_rows(min_row=3, max_row=7, min_col=1, max_col=2):
        for cell in row:
            set_excel_cell_style(cell, bold=cell.column == 1, fill_color="F9FAFB")

    # Detay sayfası başlıkları.
    headers = [
        "Soru No",
        "Soru",
        "A Şıkkı",
        "B Şıkkı",
        "C Şıkkı",
        "D Şıkkı",
        "İşaretlenen Cevap",
        "Doğru Cevap",
        "Sonuç",
    ]
    details_sheet.append(headers)

    # Başlık satırına koyu ve kurumsal görünümlü stil uygulanır.
    for cell in details_sheet[1]:
        set_excel_cell_style(cell, bold=True, fill_color="E5E7EB")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


    # Her soru sonucu detay sayfasına satır olarak eklenir.
    for index, item in enumerate(user_results, start=1):
        details_sheet.append([
            index,
            item["question"],
            item["options"]["A"],
            item["options"]["B"],
            item["options"]["C"],
            item["options"]["D"],
            item["user_answer"],
            item["correct_answer"],
            "Doğru" if item["is_correct"] else "Yanlış",
        ])

    # Detay satırlarına stil ve sonuç durumuna göre dolgu uygulanır.
    for row in details_sheet.iter_rows(min_row=2, max_row=details_sheet.max_row, min_col=1, max_col=len(headers)):
        result_cell = row[-1]
        row_fill = "DCFCE7" if result_cell.value == "Doğru" else "FEE2E2"
        for cell in row:
            set_excel_cell_style(cell, fill_color=row_fill)

    # Filtreleme ve sabit başlık özellikleri kullanıcı deneyimini artırır.
    details_sheet.auto_filter.ref = details_sheet.dimensions
    details_sheet.freeze_panes = "A2"

    # Kolon genişlikleri düzenlenir.
    autofit_excel_columns(summary_sheet, max_width=35)
    autofit_excel_columns(details_sheet, max_width=55)

    # Bazı metin yoğun kolonlarda genişlik manuel olarak daha okunabilir tutulur.
    details_sheet.column_dimensions["B"].width = 45
    details_sheet.column_dimensions["C"].width = 35
    details_sheet.column_dimensions["D"].width = 35
    details_sheet.column_dimensions["E"].width = 35
    details_sheet.column_dimensions["F"].width = 35


    # Excel dosyası belirtilen yola kaydedilir.
    workbook.save(xlsx_path)
    return xlsx_path

#############################################################################
# XML CREATE
# save_results_xml fonksiyonu:
# Quiz sonucunu XML formatında kaydeder.
# XML, sistemler arası veri taşıma ve entegrasyon için kullanışlıdır.
# Bu fonksiyon Python standart kütüphanesi ile çalışır, ekstra paket gerektirmez.
def save_results_xml(base_name, score, total, percent, user_results):
    xml_path = base_name.with_suffix(".xml")

    # Kök XML etiketi oluşturulur.
    root = ET.Element("quiz_result")

    # Özet bilgiler için summary etiketi oluşturulur.
    summary = ET.SubElement(root, "summary")
    ET.SubElement(summary, "date").text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    ET.SubElement(summary, "score").text = str(score)
    ET.SubElement(summary, "wrong").text = str(total - score)
    ET.SubElement(summary, "total").text = str(total)
    ET.SubElement(summary, "percent").text = f"{percent:.2f}"

    # Soruların tutulacağı ana questions etiketi oluşturulur.
    questions_element = ET.SubElement(root, "questions")

    # Her kullanıcı sonucu XML içine ayrı question etiketi olarak eklenir.
    for index, item in enumerate(user_results, start=1):
        question_element = ET.SubElement(questions_element, "question", {
            "no": str(index),
            "result": "correct" if item["is_correct"] else "wrong",
        })

        ET.SubElement(question_element, "text").text = item["question"]

        options_element = ET.SubElement(question_element, "options")
        for option_key, option_value in item["options"].items():
            option_element = ET.SubElement(options_element, "option", {
                "key": option_key,
                "is_user_answer": str(option_key == item["user_answer"]).lower(),
                "is_correct_answer": str(option_key == item["correct_answer"]).lower(),
            })
            option_element.text = option_value

        ET.SubElement(question_element, "user_answer").text = item["user_answer"]
        ET.SubElement(question_element, "correct_answer").text = item["correct_answer"]
        ET.SubElement(question_element, "is_correct").text = str(item["is_correct"]).lower()

    # XML çıktısı önce byte olarak üretilir.
    rough_xml = ET.tostring(root, encoding="utf-8")

    # minidom ile daha okunabilir, girintili XML çıktısı hazırlanır.
    pretty_xml = minidom.parseString(rough_xml).toprettyxml(indent="    ", encoding="utf-8")

    # XML dosyası diske yazılır.
    xml_path.write_bytes(pretty_xml)
    return xml_path


##################################################################################
# get_option_css_class fonksiyonu:
# HTML raporunda her seçeneğin hangi renkle/biçimle gösterileceğini belirler.
#
# Senaryolar:
# - Hem kullanıcı seçti hem doğruysa: özel başarılı seçili stil
# - Doğru cevap ama kullanıcı seçmediysa: doğru cevap stili
# - Kullanıcı yanlış şıkkı seçtiyse: yanlış seçili stil
# - Diğer tüm şıklar: normal stil
#
# Bu fonksiyon yalnızca CSS sınıfı adı döndürür.
def get_option_css_class(option_key, user_answer, correct_answer):
    # Kullanıcı doğru şıkkı seçtiyse bu şık hem doğru hem seçili olarak gösterilir.
    if option_key == correct_answer and option_key == user_answer:
        return "option option-correct-selected"

    # Kullanıcı yanlış cevap verdiyse doğru cevap yeşil şekilde ayrıca vurgulanır.
    if option_key == correct_answer:
        return "option option-correct"

    # Kullanıcının seçtiği ama yanlış olan şık kırmızı şekilde vurgulanır.
    if option_key == user_answer and option_key != correct_answer:
        return "option option-wrong-selected"

    # Diğer şıklar normal görünümde kalır.
    return "option"


# save_all_results fonksiyonu:
# Quiz tamamlandıktan sonra tüm rapor dosyalarını tek noktadan üretir.
# TXT, CSV, HTML, Word, Excel ve XML dosyalarını oluşturur ve bunların yolunu döndürür.
def save_all_results(score, total, user_results):
    # Başarı oranı yüzde olarak hesaplanacaktır.
    # total sıfır olursa hata olmaması için güvenli kontroller gereklidir
    percent = (score / total) * 100 if total else 0

    # Tüm dosyalarda kullanıalcak ortak temel isim üretilir.
    base_name = create_result_base_names()

    # Altı farklı formatta rapor oluşturulur.
    txt_file = save_results_txt(base_name, score, total, percent, user_results)
    csv_file = save_results_csv(base_name, score, total, percent, user_results)
    html_file = save_results_html(base_name, score, total, percent, user_results)
    docx_file = save_results_docx(base_name, score, total, percent, user_results)
    xlsx_file = save_results_xlsx(base_name, score, total, percent, user_results)
    xml_file = save_results_xml(base_name, score, total, percent, user_results)

    return txt_file, csv_file, html_file, docx_file, xlsx_file, xml_file, percent


# show_final_summary fonksiyonu:
# Quiz bitince kullanıcıya terminal ekranında özet bilgi verir.
# Doğru, yanlış, toplam ve yüzde başarı bilgisi yazdırılır.
# Ayrıca başarı oranına göre kısa bir değerlendirme metni gösterilir.
def show_final_summary(score, total, percent):
    print("\n" + "=" * 70)
    print("Quiz tamamlandı")
    print("=" * 70)
    print(f"Doğru sayısı: {score}")
    print(f"Yanlış sayısı: {total - score}")
    print(f"Toplam soru: {total}")
    print(f"Başarı Oranı: %{percent:.2f}")

    # Başarı oranına göre seviye değerlendirme
    if percent == 100:
        print("Değerlendirme mükemmel")
    elif percent >= 80:
        print("Değerlendirme çok iyi")
    elif percent >= 60:
        print("Değerlendirme iyi")
    elif percent >= 40:
        print("Değerlendirme orta")
    else:
        print("Değerlendirme: Daha iyi olabilir çok çalışmalısınız.")


################################################################################################################
# main fonksiyonu:
# Programın giriş noktasıdır.
# Genel akış burada yönetilir:
# 1) Soruları yükle
# 2) Menü göster
# 3) Kullanıcı seçimine göre quiz başlat / soru sayısını göster / çıkış yap
def main():
    # Önce CSV dosyasındaki sorular yüklenir.
    questions = load_questions(CSV_FILE)

    # Eğer soru listesi boşsa program devam etmez.
    if not questions:
        print("Program sonlandırıldı. Soru listesi boş veya CSV hatalı.")
        return

    # Sonsuz döngü ile menü sürekli gösterilir.
    # Kullanıcı '3' ile çıkış yapana kadar program açık kalır.
    while True:
        choice = prompt_menu()

        # 1 seçildiyse quiz başlatılır.
        if choice == "1":
            score, total, user_results = run_quiz(questions)
            txt_file, csv_file, html_file, docx_file, xlsx_file, xml_file, percent = save_all_results(score, total, user_results)

            show_final_summary(score, total, percent)
            print("\nSonuç dosyaları oluşturuldu:")
            print(f"TXT   : {txt_file}")
            print(f"CSV   : {csv_file}")
            print(f"HTML  : {html_file}")
            print(f"WORD  : {docx_file if docx_file else 'Oluşturulamadı - python-docx/lxml kurulumu gerekli'}")
            print(f"EXCEL : {xlsx_file if xlsx_file else 'Oluşturulamadı - openpyxl kurulumu gerekli'}")
            print(f"XML   : {xml_file}")

        # 2 seçildiyse toplam soru sayısı gösterilir.
        elif choice == "2":
            print(f"\nToplam soru sayısı: {len(questions)}")

        # 3 seçildiyse program sonlandırılır.
        elif choice == "3":
            print("Program kapatıldı.")
            break


# Bu koşulun anlamı:
# Dosya doğrudan çalıştırılırsa main() fonksiyonu devreye girsin.
# Ancak bu dosya başka bir Python dosyası içinde import edilirse
# otomatik olarak çalışmasın.
if __name__ == "__main__":
    main()
